#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Genera el Diagrama de Gantt - versión INSCRIB SYSTEM."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = r'D:\Usuarios\MCAMPOS\Desktop\INSCRIB SYSTEM\INSCRIB SYSTEM\informe de proyecto'
OUTPUT = os.path.join(OUTPUT_DIR, 'DIAGRAMA_GANTT.docx')

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.bold = True
    hs.font.size = Pt(16 - level * 2)

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_field_table(doc, left, right):
    t = doc.add_table(rows=1, cols=2)
    t.columns[0].width = Cm(4)
    t.columns[1].width = Cm(7)
    c0 = t.cell(0, 0).paragraphs[0]
    r0 = c0.add_run(left); r0.bold = True; r0.font.name = 'Times New Roman'; r0.font.size = Pt(11)
    c1 = t.cell(0, 1).paragraphs[0]
    r1 = c1.add_run(right); r1.font.name = 'Times New Roman'; r1.font.size = Pt(11)
    return t

# Título
doc.add_heading('DIAGRAMA DE GANTT', level=1)
add_para('Proyecto: Sistema de Gestión de Inscripciones y Administración Escolar "INSCRIB SYSTEM"', bold=True)
add_para('U.E. Dr. José Manuel Cova Maza - Pasantías 2026', bold=True)
add_para('Marcelo Fernando Campos Anacona | Ing. Yalitza Guevara', bold=True)
add_para('Leyenda: ■ = Actividad programada para esa semana', bold=True)
add_para('Período total: 12 semanas (480 horas) | 01 de Junio al 21 de Agosto de 2026', bold=True)

# Datos encabezado
add_field_table(doc, 'Estudiante:', 'Marcelo F. Campos Anacona')
add_field_table(doc, 'Tutor Académico:', 'Ing. Yalitza Guevara')
add_field_table(doc, 'Empresa:', 'U.E. Dr. José Manuel Cova Maza')
add_field_table(doc, 'Tutor Industrial:', 'Sol L. Romero Gil')
add_field_table(doc, 'Período:', '01 Jun - 21 Ago 2026')
add_field_table(doc, 'Duración:', '12 semanas (480 horas)')

# Tabla Gantt
headers = ['N°', 'Actividad', 'S1', 'S2', 'S3', 'S4', 'S5', 'S6', 'S7', 'S8', 'S9', 'S10', 'S11', 'S12']
rows = [
    ['1', 'Inducción y presentación', '■', '', '', '', '', '', '', '', '', '', '', ''],
    ['2', 'Conocimiento de la institución (PSA)', '■', '', '', '', '', '', '', '', '', '', '', ''],
    ['3', 'Conocimiento del Departamento de Administración Escolar', '■', '■', '', '', '', '', '', '', '', '', '', ''],
    ['4', 'Análisis de procesos de inscripción', '', '■', '■', '', '', '', '', '', '', '', '', ''],
    ['5', 'Elaboración del plan de trabajo', '', '■', '', '', '', '', '', '', '', '', '', ''],
    ['6', 'Diseño de arquitectura del sistema', '', '', '', '■', '', '', '', '', '', '', '', ''],
    ['7', 'Diseño de base de datos (SQLAlchemy)', '', '', '', '■', '', '', '', '', '', '', '', ''],
    ['8', 'Desarrollo: Autenticación y Usuarios', '', '', '', '', '■', '■', '', '', '', '', '', ''],
    ['9', 'Desarrollo: Año Escolar y Grados', '', '', '', '', '■', '', '', '', '', '', '', ''],
    ['10', 'Desarrollo: Módulo Estudiantes', '', '', '', '', '', '■', '■', '', '', '', '', ''],
    ['11', 'Desarrollo: Módulo Representantes', '', '', '', '', '', '', '■', '■', '', '', '', ''],
    ['12', 'Desarrollo: Matrícula/Inscripción', '', '', '', '', '', '', '', '■', '', '', '', ''],
    ['13', 'Desarrollo: Plantillas de documentos', '', '', '', '', '', '', '', '■', '■', '', '', ''],
    ['14', 'Desarrollo: Sitio web público', '', '', '', '', '', '', '', '', '■', '', '', ''],
    ['15', 'Desarrollo: Seguridad y Auditoría', '', '', '', '', '', '', '', '', '■', '■', '', ''],
    ['16', 'Pruebas y ajustes', '', '', '', '', '', '', '', '', '', '■', '', ''],
    ['17', 'Puesta en marcha', '', '', '', '', '', '', '', '', '', '', '■', ''],
    ['18', 'Capacitación al personal', '', '', '', '', '', '', '', '', '', '', '■', ''],
    ['19', 'Documentación y cierre', '', '', '', '', '', '', '', '', '', '', '', '■'],
    ['20', 'Soporte y seguimiento post-implementación', '', '', '', '', '', '', '', '', '', '■', '■', ''],
]

t = doc.add_table(rows=1 + len(rows), cols=len(headers))
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.style = 'Table Grid'
for i, h in enumerate(headers):
    cell = t.cell(0, i)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
for r_idx, row in enumerate(rows):
    for c_idx, val in enumerate(row):
        cell = t.cell(r_idx + 1, c_idx)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(val)); r.font.name = 'Times New Roman'; r.font.size = Pt(10)

doc.save(OUTPUT)
print(f'Documento guardado: {OUTPUT}')
print(f'Tamaño: {os.path.getsize(OUTPUT) / 1024:.1f} KB')
