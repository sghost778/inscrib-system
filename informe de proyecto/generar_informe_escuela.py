#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generador del Informe de Prácticas Profesionales - UPTJAA TSU-3
Adaptado al Sistema de Gestión de Inscripciones y Administración Escolar INSCRIB SYSTEM
desarrollado en la U.E. Dr. José Manuel Cova Maza."""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = r'D:\Usuarios\MCAMPOS\Desktop\INSCRIB SYSTEM\INSCRIB SYSTEM\informe de proyecto'
OUTPUT = os.path.join(OUTPUT_DIR, 'INFORME_PASANTIAS_MARCELO_CAMPOS.docx')

doc = Document()

# ─── Estilos ───────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    hs.font.size = Pt(16 - level * 2)

# ─── Helper functions ──────────────────────────────────────────
def add_field_table(doc, left_label, right_value):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.columns[0].width = Cm(5)
    t.columns[1].width = Cm(12)
    cell0 = t.cell(0, 0)
    cell1 = t.cell(0, 1)
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run0 = p0.add_run(left_label)
    run0.bold = True
    run0.font.name = 'Times New Roman'
    run0.font.size = Pt(12)
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p1.add_run(str(right_value))
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    for row in t.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    return t

def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5 + level * 1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    return t

# ══════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para(doc, 'REPÚBLICA BOLIVARIANA DE VENEZUELA', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, 'UNIVERSIDAD POLITÉCNICA TERRITORIAL', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, '"JOSÉ ANTONIO ANZOÁTEGUI"', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, 'Programa Nacional de Formación en Ingeniería de Sistemas', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, 'Trayecto 3 - Fase 1', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()

add_para(doc, 'DESARROLLO DE UN SISTEMA WEB DE GESTIÓN DE', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, 'INSCRIPCIONES Y ADMINISTRACIÓN ESCOLAR', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, '"INSCRIB SYSTEM" EN LA U.E. DR. JOSÉ MANUEL COVA MAZA', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
doc.add_paragraph()

add_para(doc, 'INFORME DE PRÁCTICAS PROFESIONALES', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()

add_para(doc, 'Autor: Marcelo Fernando Campos Anacona', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, 'C.I: V-32.062.637', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()
add_para(doc, 'Tutora Académica: Ing. Yalitza Guevara', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, 'Tutora Industrial: Sol Liliana Romero Gil', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()
add_para(doc, 'Puerto Ordaz, Agosto 2026', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CONTRAPORTADA
# ══════════════════════════════════════════════════════════════
for _ in range(10):
    doc.add_paragraph()
add_para(doc, 'DESARROLLO DE UN SISTEMA WEB DE GESTIÓN DE', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, 'INSCRIPCIONES Y ADMINISTRACIÓN ESCOLAR', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para(doc, '"INSCRIB SYSTEM" EN LA U.E. DR. JOSÉ MANUEL COVA MAZA', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
doc.add_paragraph()
add_para(doc, 'Marcelo Fernando Campos Anacona', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()
add_para(doc, 'Informe de Prácticas Profesionales presentado como requisito parcial', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, 'para aprobar la unidad curricular Pasantías Profesionales del', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, 'Programa Nacional de Formación en Ingeniería de Sistemas', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, 'Trayecto 3 - Fase 1', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()
add_para(doc, 'Tutora Académica: Ing. Yalitza Guevara', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, 'Tutora Industrial: Sol Liliana Romero Gil', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
doc.add_paragraph()
add_para(doc, 'Puerto Ordaz, Agosto 2026', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# EVALUACIÓN DEL TUTOR INDUSTRIAL (blank form)
# ══════════════════════════════════════════════════════════════
doc.add_heading('EVALUACIÓN DEL TUTOR INDUSTRIAL', level=1)
add_para(doc, '(Página para ser llenada por el Tutor Industrial)')
doc.add_paragraph()
add_para(doc, 'Se deja esta página para que el Tutor Industrial registre su evaluación del desempeño del pasante, según el formato institucional de la UPTJAA.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# EVALUACIÓN DEL TUTOR ACADÉMICO (blank form)
# ══════════════════════════════════════════════════════════════
doc.add_heading('EVALUACIÓN DEL TUTOR ACADÉMICO', level=1)
add_para(doc, '(Página para ser llenada por el Tutor Académico)')
doc.add_paragraph()
add_para(doc, 'Se deja esta página para que el Tutor Académico registre su evaluación del informe de pasantías, según el formato institucional de la UPTJAA.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# DEDICATORIA
# ══════════════════════════════════════════════════════════════
doc.add_heading('DEDICATORIA', level=1)
doc.add_paragraph()
add_para(doc, 'A Dios, por ser mi guía y fortaleza en cada paso de este camino.')
add_para(doc, 'A mis padres, por su amor incondicional, su sacrificio y por enseñarme que con perseverancia todo es posible. Este logro es tanto mío como suyo.')
add_para(doc, 'A mi familia, por su apoyo constante y por creer en mí incluso en los momentos más difíciles.')
add_para(doc, 'A mis profesores, por compartir su conocimiento y por formarme como profesional y como persona.')
add_para(doc, 'A todo aquel que de una u otra manera contribuyó a la realización de este proyecto.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# AGRADECIMIENTO
# ══════════════════════════════════════════════════════════════
doc.add_heading('AGRADECIMIENTO', level=1)
doc.add_paragraph()
agradecimientos = [
    'A Dios por la salud, sabiduría y fortaleza para alcanzar esta meta.',
    'A mi familia por su apoyo incondicional, paciencia y comprensión durante todo este proceso.',
    'A la Universidad Politécnica Territorial "José Antonio Anzoátegui" (UPTJAA) y a todos sus profesores por la formación académica recibida.',
    'A la Unidad Educativa Dr. José Manuel Cova Maza por abrirme las puertas y brindarme la oportunidad de desarrollar mis prácticas profesionales en sus instalaciones.',
    'A mi tutora industrial, Sol Liliana Romero Gil, Coordinadora Académica de la institución escolar, por su orientación, confianza y valiosas enseñanzas durante todo el período de pasantías.',
    'A mi tutora académica, Ing. Yalitza Guevara, por su seguimiento, asesoría y dedicación en la supervisión de este proyecto.',
    'A todo el personal docente y administrativo de la institución, por su calidez, colaboración y disposición para compartir sus conocimientos.',
    'A mis compañeros de estudio y amigos, por su apoyo moral y por hacer de esta etapa universitaria una experiencia inolvidable.',
]
for a in agradecimientos:
    add_para(doc, a)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ÍNDICE GENERAL (placeholder)
# ══════════════════════════════════════════════════════════════
doc.add_heading('ÍNDICE GENERAL', level=1)
add_para(doc, '(Nota: El índice general debe generarse automáticamente al finalizar la edición del documento en Word usando la función de tabla de contenido).')
doc.add_paragraph()
add_para(doc, 'CONTENIDO')
toc_items = [
    ('PORTADA', 'i'),
    ('CONTRAPORTADA', 'ii'),
    ('EVALUACIÓN DEL TUTOR INDUSTRIAL', 'iii'),
    ('EVALUACIÓN DEL TUTOR ACADÉMICO', 'iv'),
    ('DEDICATORIA', 'v'),
    ('AGRADECIMIENTO', 'vi'),
    ('ÍNDICE GENERAL', 'vii'),
    ('ÍNDICE DE TABLAS', 'viii'),
    ('ÍNDICE DE FIGURAS', 'ix'),
    ('RESUMEN', 'x'),
    ('INTRODUCCIÓN', '1'),
    ('CAPÍTULO I: DESCRIPCIÓN DE LA ORGANIZACIÓN', '3'),
    ('CAPÍTULO II: DESCRIPCIÓN DE LA SITUACIÓN ORGANIZACIONAL', '8'),
    ('CAPÍTULO III: ACTIVIDADES EJECUTADAS Y RESULTADOS', '14'),
    ('CAPÍTULO IV: CONCLUSIONES Y RECOMENDACIONES', '20'),
    ('REFERENCIAS', '23'),
    ('ANEXOS', '24'),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title} .......................................................... {page}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ÍNDICE DE TABLAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('ÍNDICE DE TABLAS', level=1)
table_of_tables = [
    ('Tabla 1', 'Datos de la Organización', '7'),
    ('Tabla 2', 'Cronograma de Actividades (Diagrama de Gantt)', '13'),
    ('Tabla 3', 'Actividades Ejecutadas Semana 1', '15'),
    ('Tabla 4', 'Actividades Ejecutadas Semana 2', '15'),
    ('Tabla 5', 'Actividades Ejecutadas Semana 3', '16'),
    ('Tabla 6', 'Actividades Ejecutadas Semana 4', '16'),
    ('Tabla 7', 'Actividades Ejecutadas Semana 5', '16'),
    ('Tabla 8', 'Actividades Ejecutadas Semana 6', '17'),
    ('Tabla 9', 'Actividades Ejecutadas Semana 7', '17'),
    ('Tabla 10', 'Actividades Ejecutadas Semana 8', '17'),
    ('Tabla 11', 'Actividades Ejecutadas Semana 9', '18'),
    ('Tabla 12', 'Actividades Ejecutadas Semana 10', '18'),
    ('Tabla 13', 'Actividades Ejecutadas Semana 11', '18'),
    ('Tabla 14', 'Actividades Ejecutadas Semana 12', '19'),
    ('Tabla 15', 'Resultados Cuantitativos del Sistema', '19'),
]
for num, desc, page in table_of_tables:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}: {desc} .............................................. {page}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ÍNDICE DE FIGURAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('ÍNDICE DE FIGURAS', level=1)
figure_list = [
    ('Figura 1', 'Ubicación Geográfica de la U.E. Dr. José Manuel Cova Maza', '5'),
    ('Figura 2', 'Organigrama General de la Institución', '6'),
    ('Figura 3', 'Estructura del Departamento de Administración Escolar / Informática', '7'),
    ('Figura 4', 'Arquitectura del Sistema INSCRIB SYSTEM', '11'),
    ('Figura 5', 'Dashboard Principal del Panel Administrativo', '19'),
]
for num, desc, page in figure_list:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}: {desc} .............................................. {page}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# RESUMEN
# ══════════════════════════════════════════════════════════════
doc.add_heading('RESUMEN', level=1)
doc.add_paragraph()
resumen_text = (
    'El presente informe describe el desarrollo de un Sistema de Gestión de Inscripciones y '
    'Administración Escolar denominado "INSCRIB SYSTEM", realizado como parte de las prácticas '
    'profesionales del Programa Nacional de Formación en Ingeniería de Sistemas de la UPTJAA, en la '
    'Unidad Educativa Dr. José Manuel Cova Maza. El proyecto tuvo como objetivo principal automatizar '
    'los procesos de inscripción y matrícula escolar, la gestión de estudiantes, representantes y la '
    'administración del sitio web institucional público. Se desarrolló una aplicación web utilizando '
    'el microframework Flask (Python 3.13), SQLAlchemy, Flask-SQLAlchemy, Flask-CORS, Flask-Talisman '
    '(cabeceras de seguridad y CSP), Flask-Limiter (limitación de peticiones), bcrypt (hash de '
    'contraseñas) y python-docx, con base de datos SQLite y un frontend HTML5, CSS3, JavaScript, '
    'Font Awesome y PWA. El sistema cuenta con módulos de autenticación, gestión de usuarios con roles '
    '(administrador y secretario), año escolar, estudiantes, representantes, matrícula/inscripción, '
    'plantillas de documentos y administración del sitio web, así como un portal público para los '
    'representantes. Como resultado se logró la digitalización de los procesos de admisión, la '
    'trazabilidad mediante una tabla de auditoría, el fortalecimiento de la seguridad y la publicación '
    'de la información institucional en la web.'
)
add_para(doc, resumen_text)
doc.add_paragraph()
add_para(doc, 'Palabras clave: Sistema de Inscripciones, Administración Escolar, Flask, SQLAlchemy, '
           'Seguridad Web, Auditoría, PWA, Trazabilidad.',
           align=WD_ALIGN_PARAGRAPH.LEFT, size=11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════
doc.add_heading('INTRODUCCIÓN', level=1)
doc.add_paragraph()
intro_paras = [
    'La gestión eficiente de los procesos de inscripción, matrícula y administración escolar es un '
    'factor crítico para las instituciones educativas. La capacidad de mantener un control preciso de '
    'los estudiantes, sus representantes, los años escolares y la documentación asociada, así como '
    'publicar información institucional actualizada, son elementos fundamentales para la calidad '
    'educativa y la transparencia de la organización.',

    'La Unidad Educativa Dr. José Manuel Cova Maza, institución que ofrece niveles de Educación '
    'Inicial, Primaria, Media General (Bachillerato), Educación Especial y Formación Docente, '
    'identificó la necesidad de automatizar sus procesos de inscripción y administración escolar para '
    'superar las limitaciones del sistema manual que venía operando, el cual presentaba registros en '
    'hojas de cálculo, falta de trazabilidad y demoras en la atención a representantes.',

    'En este contexto, se planteó el desarrollo del Sistema de Gestión de Inscripciones y '
    'Administración Escolar "INSCRIB SYSTEM", una aplicación web construida con Flask que permite '
    'gestionar inscripciones, estudiantes, representantes, matrículas y el sitio web institucional. '
    'El proyecto fue ejecutado durante un período de 12 semanas (480 horas), comprendido entre el 01 '
    'de Junio y el 21 de Agosto de 2026.',

    'El presente informe está estructurado en cuatro capítulos. El Capítulo I describe la organización '
    'donde se desarrollaron las prácticas profesionales. El Capítulo II presenta la situación '
    'organizacional identificada, los objetivos del proyecto y las actividades planificadas. El '
    'Capítulo III detalla las actividades ejecutadas y los resultados obtenidos. Finalmente, el '
    'Capítulo IV expone las conclusiones y recomendaciones derivadas del proyecto.',
]
for p_text in intro_paras:
    add_para(doc, p_text)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CAPÍTULO I: DESCRIPCIÓN DE LA ORGANIZACIÓN
# ══════════════════════════════════════════════════════════════
doc.add_heading('CAPÍTULO I', level=1)
doc.add_heading('DESCRIPCIÓN DE LA ORGANIZACIÓN', level=2)

# 1.1 Nombre de la Organización
doc.add_heading('1.1. Nombre de la Organización', level=3)
add_para(doc, 'La organización donde se desarrollaron las prácticas profesionales es la Unidad '
           'Educativa Dr. José Manuel Cova Maza (U.E. Dr. José Manuel Cova Maza), institución '
           'educativa venezolana ubicada en Puerto Ordaz, Estado Bolívar.')

# 1.2 Razón Social
doc.add_heading('1.2. Razón Social', level=3)
add_field_table(doc, 'Razón Social:', 'U.E. Dr. José Manuel Cova Maza')
add_field_table(doc, 'Dirección:', 'Av. Principal, Puerto Ordaz, Estado Bolívar')
add_field_table(doc, 'Teléfono:', '+58 412-1234567')
add_field_table(doc, 'Correo electrónico:', 'info@uejmcm.edu.ve')
add_field_table(doc, 'Sector:', 'Educación')
add_field_table(doc, 'Niveles que ofrece:', 'Educación Inicial, Primaria, Media General (Bachillerato), Educación Especial y Formación Docente')

# 1.3 Reseña Histórica
doc.add_heading('1.3. Reseña Histórica de la Organización', level=3)
add_para(doc, 'La Unidad Educativa Dr. José Manuel Cova Maza es una institución educativa de la '
           'localidad de Puerto Ordaz, Estado Bolívar, cuyo nombre honra la memoria del Dr. José '
           'Manuel Cova Maza, destacada figura de la educación regional. A lo largo de su trayectoria, '
           'la institución ha formado generaciones de estudiantes en los distintos niveles del sistema '
           'educativo venezolano.')
add_para(doc, 'La institución ha evolucionado desde una oferta académica básica hasta abarcar la '
           'Educación Inicial, Primaria, Media General (Bachillerato), Educación Especial y la '
           'Formación Docente, adaptándose a las necesidades de la comunidad y a las políticas '
           'educativas nacionales. En los últimos años, la dirección de la institución ha impulsado '
           'la modernización de sus procesos administrativos mediante el uso de herramientas '
           'tecnológicas, lo que motivó la implementación del sistema INSCRIB SYSTEM.')

# 1.4 Ubicación Geográfica
doc.add_heading('1.4. Ubicación Geográfica de la Organización', level=3)
add_para(doc, 'La institución se encuentra en la Avenida Principal de Puerto Ordaz, Estado Bolívar, '
           'Venezuela. Esta ubicación permite el acceso de la comunidad estudiantil y de los '
           'representantes desde diversos sectores de la ciudad y de la región guayanesa.')
add_para(doc, '(Ver ubicación en Anexos - Mapa de Ubicación)')

# 1.5 Visión
doc.add_heading('1.5. Visión de la Organización', level=3)
add_para(doc, '"Ser una institución educativa de excelencia, formadora de ciudadanos íntegros y '
           'comprometidos con el desarrollo de la región y del país, reconocida por la calidad de su '
           'educación, la innovación de sus procesos y el uso responsable de la tecnología."')

# 1.6 Misión
doc.add_heading('1.6. Misión de la Organización', level=3)
add_para(doc, '"Brindar una formación integral y de calidad en los niveles de Educación Inicial, '
           'Primaria, Media General, Educación Especial y Formación Docente, fundamentada en valores '
           'humanos, el pensamiento crítico y el aprovechamiento de las herramientas tecnológicas para '
           'el beneficio de la comunidad."')

# 1.7 Organigrama
doc.add_heading('1.7. Organigrama de la Organización', level=3)
add_para(doc, 'La estructura organizativa de la U.E. Dr. José Manuel Cova Maza se compone de los '
           'siguientes niveles jerárquicos:')
add_bullet(doc, 'Dirección: máxima autoridad de la institución, responsable de la gestión académica y administrativa general.')
add_bullet(doc, 'Coordinación Académica: a cargo de la planificación curricular, horarios y seguimiento del rendimiento escolar.')
add_bullet(doc, 'Secretaría: encargada de la gestión documental, inscripciones, matrículas y archivo de expedientes.')
add_bullet(doc, 'Docencia: cuerpo de docentes por nivel (Inicial, Primaria, Media General, Especial y Formación Docente).')
add_bullet(doc, 'Administración: recursos humanos, finanzas, mantenimiento e infraestructura.')
add_para(doc, '(Ver organigrama detallado en Anexos)')

# 1.8 Estructura Organizativa
doc.add_heading('1.8. Estructura Organizativa', level=3)
add_para(doc, 'La institución opera bajo un esquema jerárquico donde la Dirección ejerce la rectoría '
           'de la institución y coordina con la Coordinación Académica y la Secretaría para la '
           'operatividad diaria. Los departamentos de Docencia y Administración reportan a la '
           'Dirección y se articulan con la Coordinación Académica para el cumplimiento del calendario '
           'escolar y de los objetivos institucionales.')

# 1.9 Unidad / Departamento
doc.add_heading('1.9. Estructura Organizativa de la Unidad o Departamento donde se Desarrolló la Práctica Profesional', level=3)
add_para(doc, 'Las prácticas profesionales se desarrollaron en el Departamento de Administración '
           'Escolar / Informática, bajo la supervisión de la Coordinadora Académica, Sol Liliana '
           'Romero Gil. Este departamento es el encargado de gestionar las inscripciones, matrículas, '
           'el registro de estudiantes y representantes, así como la administración de la '
           'infraestructura tecnológica y del sitio web institucional.')
add_para(doc, 'El departamento está conformado por:')
add_bullet(doc, 'Coordinador Académico / Coordinadora Académica')
add_bullet(doc, 'Secretario(a) de la institución')
add_bullet(doc, 'Personal de informática / soporte técnico')
add_bullet(doc, 'Asistentes administrativos de inscripción y archivo')

# 1.10 Descripción actividad productiva
doc.add_heading('1.10. Descripción de la Actividad Productiva', level=3)
add_para(doc, 'La actividad principal de la U.E. Dr. José Manuel Cova Maza es la prestación del '
           'servicio educativo en sus distintos niveles, lo que incluye:')
add_bullet(doc, 'Educación Inicial: atención y estimulación temprana de niños y niñas.')
add_bullet(doc, 'Educación Primaria: formación básica y desarrollo de competencias.')
add_bullet(doc, 'Media General (Bachillerato): preparación académica para la educación superior.')
add_bullet(doc, 'Educación Especial: atención a la diversidad y necesidades educativas especiales.')
add_bullet(doc, 'Formación Docente: capacitación y actualización de personal educativo.')
add_bullet(doc, 'Gestión administrativa escolar: inscripciones, matrículas, control de estudiantes y representantes.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CAPÍTULO II: DESCRIPCIÓN DE LA SITUACIÓN ORGANIZACIONAL
# ══════════════════════════════════════════════════════════════
doc.add_heading('CAPÍTULO II', level=1)
doc.add_heading('DESCRIPCIÓN DE LA SITUACIÓN ORGANIZACIONAL', level=2)

# 2.1 Situación Organizacional
doc.add_heading('2.1. Situación Organizacional', level=3)
add_para(doc, 'El departamento de Administración Escolar de la U.E. Dr. José Manuel Cova Maza '
           'enfrentaba diversos desafíos en la gestión de sus procesos de inscripción y matrícula. '
           'Los procesos se realizaban de forma manual, utilizando hojas de cálculo y registros '
           'físicos, lo que generaba:')
add_bullet(doc, 'Tiempos prolongados en los procesos de inscripción y matrícula escolar.')
add_bullet(doc, 'Falta de trazabilidad de las modificaciones sobre los expedientes de estudiantes y representantes.')
add_bullet(doc, 'Dificultades para consultar rápidamente la información de un estudiante o representante.')
add_bullet(doc, 'Errores e inconsistencias en los registros mantenidos en hojas de cálculo.')
add_bullet(doc, 'Demoras en la generación de documentos y reportes para la dirección y los representantes.')
add_bullet(doc, 'Ausencia de control de acceso por roles y de respaldo automatizado de la información.')
add_bullet(doc, 'Información institucional (noticias, programas, galería) desactualizada o no publicada en la web.')

# 2.2 Práctica Profesional
doc.add_heading('2.2. Práctica Profesional', level=3)
add_para(doc, 'Ante la situación descrita, se propuso el desarrollo del Sistema de Gestión de '
           'Inscripciones y Administración Escolar "INSCRIB SYSTEM" como solución tecnológica para '
           'optimizar los procesos administrativos de la institución. Las prácticas profesionales se '
           'enfocaron en el análisis, diseño, desarrollo e implementación de este sistema, aplicando '
           'los conocimientos adquiridos durante la formación académica en Ingeniería de Sistemas.')

# 2.3 Propósito
doc.add_heading('2.3. Propósito de la Práctica Profesional', level=3)
add_para(doc, 'El propósito de la práctica profesional fue aplicar los conocimientos teórico-'
           'prácticos adquiridos durante la formación universitaria en un entorno laboral real, '
           'contribuyendo a la solución de problemas concretos de la organización y desarrollando '
           'competencias profesionales en el área de desarrollo de sistemas de información para la '
           'gestión educativa.')

# 2.4 Objetivo General
doc.add_heading('2.4. Objetivo General', level=3)
add_para(doc, 'Desarrollar e implementar un sistema web de gestión de inscripciones y administración '
           'escolar que permita automatizar los procesos de matrícula, la gestión de estudiantes y '
           'representantes y la administración del sitio web institucional en la U.E. Dr. José Manuel '
           'Cova Maza.', bold=True)

# 2.5 Objetivos Específicos
doc.add_heading('2.5. Objetivos Específicos', level=3)
objetivos = [
    'Diagnosticar los procesos actuales de inscripción y administración escolar del departamento.',
    'Diseñar la arquitectura del sistema utilizando el framework Flask y tecnologías web modernas.',
    'Desarrollar los módulos de autenticación, gestión de usuarios y año escolar.',
    'Implementar los módulos de estudiantes, representantes y matrícula/inscripción.',
    'Desarrollar el sitio web público institucional (noticias, programas, galería, contacto).',
    'Implementar controles de seguridad (bcrypt, roles, Flask-Talisman, Flask-Limiter) y registro de auditoría.',
    'Poner en marcha el sistema y capacitar al personal en su uso.',
]
for i, obj in enumerate(objetivos, 1):
    add_para(doc, f'{i}. {obj}')

# 2.6 Actividades Planificadas
doc.add_heading('2.6. Actividades Planificadas', level=3)
add_para(doc, 'Las actividades planificadas para el desarrollo del proyecto se organizaron en 12 '
           'semanas, distribuidas de la siguiente manera:')
add_bullet(doc, 'Semana 1-2: Inducción en la institución, conocimiento del departamento y elaboración del plan de trabajo.')
add_bullet(doc, 'Semana 3-4: Análisis de procesos, diseño de arquitectura y diseño de base de datos.')
add_bullet(doc, 'Semana 5-8: Desarrollo de módulos (Estudiantes, Representantes, Matrícula/Inscripción).')
add_bullet(doc, 'Semana 9-10: Seguridad, auditoría y desarrollo del sitio web público.')
add_bullet(doc, 'Semana 11-12: Pruebas, puesta en marcha, capacitación y cierre.')

# 2.7 Cronograma
doc.add_heading('2.7. Cronograma de Actividades', level=3)
add_para(doc, 'El siguiente cuadro presenta el cronograma de actividades (Diagrama de Gantt) '
           'ejecutado durante las 12 semanas de pasantías:')
gantt_headers = ['N°', 'Actividad', 'S1', 'S2', 'S3', 'S4', 'S5', 'S6', 'S7', 'S8', 'S9', 'S10', 'S11', 'S12']
gantt_rows = [
    ['1', 'Inducción y presentación', 'X', '', '', '', '', '', '', '', '', '', '', ''],
    ['2', 'Conocimiento de la institución (PSA)', 'X', '', '', '', '', '', '', '', '', '', '', ''],
    ['3', 'Conocimiento del Departamento de Administración Escolar', 'X', 'X', '', '', '', '', '', '', '', '', '', ''],
    ['4', 'Análisis de procesos de inscripción', '', 'X', 'X', '', '', '', '', '', '', '', '', ''],
    ['5', 'Elaboración del plan de trabajo', '', 'X', '', '', '', '', '', '', '', '', '', ''],
    ['6', 'Diseño de arquitectura Flask', '', '', '', 'X', '', '', '', '', '', '', '', ''],
    ['7', 'Diseño de base de datos (SQLAlchemy)', '', '', '', 'X', '', '', '', '', '', '', '', ''],
    ['8', 'Desarrollo: Autenticación y Usuarios', '', '', '', '', 'X', 'X', '', '', '', '', '', ''],
    ['9', 'Desarrollo: Año Escolar y Grados', '', '', '', '', 'X', '', '', '', '', '', '', ''],
    ['10', 'Desarrollo: Módulo Estudiantes', '', '', '', '', '', 'X', 'X', '', '', '', '', ''],
    ['11', 'Desarrollo: Módulo Representantes', '', '', '', '', '', '', 'X', 'X', '', '', '', ''],
    ['12', 'Desarrollo: Matrícula/Inscripción', '', '', '', '', '', '', '', 'X', '', '', '', ''],
    ['13', 'Desarrollo: Plantillas de documentos', '', '', '', '', '', '', '', 'X', 'X', '', '', ''],
    ['14', 'Desarrollo: Sitio web público', '', '', '', '', '', '', '', '', 'X', '', '', ''],
    ['15', 'Desarrollo: Seguridad y Auditoría', '', '', '', '', '', '', '', '', 'X', 'X', '', ''],
    ['16', 'Pruebas y ajustes', '', '', '', '', '', '', '', '', '', 'X', '', ''],
    ['17', 'Puesta en marcha', '', '', '', '', '', '', '', '', '', '', 'X', ''],
    ['18', 'Capacitación al personal', '', '', '', '', '', '', '', '', '', '', 'X', ''],
    ['19', 'Documentación y cierre', '', '', '', '', '', '', '', '', '', '', '', 'X'],
]
add_table(doc, gantt_headers, gantt_rows)

# 2.8 Justificación
doc.add_heading('2.8. Justificación', level=3)
add_para(doc, 'El proyecto se justifica desde múltiples perspectivas:')
add_para(doc, 'Desde el punto de vista operativo, la automatización de los procesos de inscripción y '
           'matrícula permite reducir los tiempos de atención a los representantes, eliminar errores '
           'de registro y garantizar la trazabilidad de los expedientes escolares.')
add_para(doc, 'Desde la perspectiva tecnológica, la implementación de un sistema con Flask, '
           'SQLAlchemy y herramientas open source demuestra la viabilidad de desarrollar soluciones '
           'escolares robustas y escalables sin depender de licencias costosas.')
add_para(doc, 'En el ámbito académico, el proyecto permite aplicar los conocimientos adquiridos '
           'durante la carrera de Ingeniería de Sistemas en un entorno real, integrando áreas como '
           'bases de datos, desarrollo web, seguridad informática y gestión de proyectos.')
add_para(doc, 'Finalmente, desde el punto de vista social, el sistema contribuye a mejorar la '
           'eficiencia de una institución educativa, impactando positivamente en la comunidad.')

# 2.9 Factibilidad
doc.add_heading('2.9. Factibilidad de Ejecución', level=3)
add_para(doc, 'El proyecto fue factible gracias a los siguientes factores:')
add_para(doc, 'Factibilidad Técnica: Se contó con los recursos tecnológicos necesarios (equipos de '
           'cómputo, software libre) y los conocimientos técnicos para el desarrollo del sistema.')
add_para(doc, 'Factibilidad Operativa: La institución mostró disposición y apoyo para la '
           'implementación del sistema, asignando una tutora industrial y facilitando el acceso a la '
           'información y a los procesos.')
add_para(doc, 'Factibilidad Económica: El uso de tecnologías open source eliminó los costos de '
           'licencias y los recursos de hardware estaban disponibles en la institución.')
add_para(doc, 'Factibilidad Temporal: El período de 12 semanas (480 horas) fue adecuado para cumplir '
           'con los objetivos planteados, siguiendo una planificación estructurada.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CAPÍTULO III: ACTIVIDADES EJECUTADAS Y RESULTADOS
# ══════════════════════════════════════════════════════════════
doc.add_heading('CAPÍTULO III', level=1)
doc.add_heading('ACTIVIDADES EJECUTADAS Y RESULTADOS', level=2)

# 3.1 Descripción
doc.add_heading('3.1. Descripción de las Actividades Ejecutadas', level=3)
add_para(doc, 'A continuación se presentan las actividades ejecutadas durante las 12 semanas de '
           'prácticas profesionales, detalladas por semana:')

weekly_activities = [
    ('Semana 1 (01 al 05 de Junio)', [
        'Presentación a la tutora industrial y al personal involucrado en la pasantía.',
        'Recorrido por las instalaciones de la institución para conocer las áreas administrativas.',
        'Participación en charlas de higiene, seguridad y ambiente (PSA) de la institución.',
        'Inducción al Departamento de Administración Escolar / Informática.',
        'Conocimiento de los procesos de inscripción, matrícula y archivo de expedientes.',
        'Revisión de procedimientos administrativos y documentación de procesos actuales.',
        'Asignación del tema de pasantías y definición del alcance del proyecto.',
    ]),
    ('Semana 2 (08 al 12 de Junio)', [
        'Elaboración del plan de trabajo en conjunto con la tutora industrial.',
        'Culminación del plan de trabajo y definición de cronograma y actividades.',
        'Revisión de las hojas de cálculo utilizadas para el control de estudiantes.',
        'Inducción sobre los niveles educativos ofrecidos por la institución.',
        'Identificación de las necesidades de información de los representantes.',
    ]),
    ('Semana 3 (15 al 19 de Junio)', [
        'Análisis de los procesos de inscripción y matrícula escolar.',
        'Levantamiento de requerimientos funcionales del sistema.',
        'Análisis del flujo de información entre Secretaría, Coordinación y Docencia.',
        'Identificación de los actores (administrador, secretario, representante).',
        'Documentación de especificaciones para el sistema INSCRIB SYSTEM.',
    ]),
    ('Semana 4 (22 al 26 de Junio)', [
        'Diseño de la arquitectura del sistema con Flask y SQLAlchemy.',
        'Selección de tecnologías (Flask-CORS, Flask-Talisman, Flask-Limiter, bcrypt, PWA).',
        'Diseño de la base de datos y modelamiento de entidades.',
        'Definición de roles y política de seguridad.',
        'Configuración del entorno de desarrollo y del servidor Flask.',
    ]),
    ('Semana 5 (29 de Junio al 03 de Julio)', [
        'Desarrollo del módulo de Autenticación (login seguro con bcrypt).',
        'Implementación del modelo de datos para el Año Escolar y Grados.',
        'Desarrollo de la gestión de Usuarios con roles (admin y secretario).',
        'Configuración de Flask-Talisman (cabeceras de seguridad y CSP).',
        'Pruebas de autenticación y control de acceso.',
    ]),
    ('Semana 6 (06 al 10 de Julio)', [
        'Desarrollo del módulo de Estudiantes (registro y listado).',
        'Implementación de formularios de registro con validaciones.',
        'Desarrollo de la consulta de estudiantes por año escolar y grado.',
        'Integración con el módulo de Representantes.',
        'Pruebas del módulo de estudiantes y ajustes de funcionalidad.',
    ]),
    ('Semana 7 (13 al 17 de Julio)', [
        'Desarrollo del módulo de Representantes (registro y listado).',
        'Implementación de la relación familiar Representante-Estudiante.',
        'Desarrollo de la consulta de representantes y estudiantes asociados.',
        'Generación de documentos mediante plantillas (python-docx).',
        'Pruebas del módulo de representantes.',
    ]),
    ('Semana 8 (20 al 24 de Julio)', [
        'Desarrollo del módulo de Matrícula/Inscripción.',
        'Implementación de las columnas de control de la inscripción (estado, fecha_retiro, lapso_registro, motivo_retiro).',
        'Desarrollo de la asociación inscripción-estudiante-año escolar-grado.',
        'Implementación de reportes de matrícula por grado.',
        'Pruebas del módulo de matrícula e inscripción.',
    ]),
    ('Semana 9 (27 al 31 de Julio)', [
        'Desarrollo del sitio web público (Inicio, Nosotros, Programas).',
        'Implementación de Noticias, Galería y Formulario de Contacto.',
        'Desarrollo del Portal del Representante (registro y consulta).',
        'Desarrollo del módulo de Seguridad y Auditoría (REGISTRO_AUDITORIA).',
        'Pruebas integrales del sistema.',
    ]),
    ('Semana 10 (03 al 07 de Agosto)', [
        'Implementación de Flask-Limiter (limitación de peticiones/rate limiting).',
        'Pruebas de funcionamiento del sistema completo (admin y sitio público).',
        'Corrección de errores identificados en las pruebas.',
        'Optimización de consultas y ajustes de la interfaz de usuario.',
        'Documentación técnica del sistema.',
    ]),
    ('Semana 11 (10 al 14 de Agosto)', [
        'Puesta en marcha del sistema (admin puerto 5001, sitio público puerto 5002).',
        'Capacitación al personal de la Secretaría y Coordinación Académica.',
        'Capacitación a usuarios del sistema (demostración práctica).',
        'Resolución de dudas y soporte técnico post-implementación.',
        'Documentación de usuario y manuales.',
    ]),
    ('Semana 12 (17 al 21 de Agosto)', [
        'Revisión final del sistema y ajustes menores.',
        'Elaboración del informe de pasantías.',
        'Inclusión de resultados y análisis.',
        'Revisión y corrección del informe.',
        'Entrega del informe y cierre de pasantías.',
    ]),
]

for week_title, activities in weekly_activities:
    doc.add_heading(week_title, level=4)
    for act in activities:
        add_bullet(doc, act)

# 3.2 Resultados
doc.add_heading('3.2. Resultados', level=3)
add_para(doc, 'Como resultado del proyecto, se logró desarrollar e implementar exitosamente el '
           'Sistema de Gestión de Inscripciones y Administración Escolar "INSCRIB SYSTEM" en la '
           'U.E. Dr. José Manuel Cova Maza, obteniendo los siguientes resultados cuantitativos y '
           'cualitativos:')

doc.add_heading('Resultados Cuantitativos', level=4)
add_table(doc, ['Indicador', 'Valor'], [
    ['Módulos funcionales implementados', '9'],
    ['Tablas de base de datos', '20'],
    ['Usuarios por defecto (admin/secretario)', '2'],
    ['Grados precargados (1er a 5to Año)', '5'],
    ['Horas de desarrollo y ejecución', '480'],
    ['Semanas de ejecución', '12'],
    ['Usuarios capacitados', '5+'],
])

doc.add_heading('Módulos del Sistema', level=4)
add_table(doc, ['N°', 'Módulo', 'Descripción'], [
    ['1', 'Autenticación / Login', 'Acceso seguro con hash bcrypt y control de sesiones'],
    ['2', 'Gestión de Usuarios', 'Roles administrador y secretario'],
    ['3', 'Año Escolar', 'Gestión y activación del período escolar'],
    ['4', 'Estudiantes', 'Registro, listado y consulta de estudiantes'],
    ['5', 'Representantes', 'Registro, listado y consulta de representantes'],
    ['6', 'Matrícula / Inscripción', 'Asociación estudiante-año-grado y control de retiro'],
    ['7', 'Plantillas de Documentos', 'Generación de documentos con python-docx'],
    ['8', 'Administración del Sitio Web', 'Noticias, Programas, Galería y Configuración'],
    ['9', 'Seguridad y Auditoría', 'Talisman, Limiter, roles y REGISTRO_AUDITORIA'],
])

doc.add_heading('Tecnologías Utilizadas', level=4)
add_table(doc, ['Tecnología', 'Versión', 'Propósito'], [
    ['Flask', '3.1.1', 'Microframework web backend (Python 3.13)'],
    ['Flask-SQLAlchemy', '3.1.1', 'ORM sobre SQLAlchemy 2.0'],
    ['SQLAlchemy', '2.0', 'Mapeo objeto-relacional'],
    ['Flask-CORS', '-', 'Gestión de políticas de origen cruzado'],
    ['Flask-Talisman', '-', 'Cabeceras de seguridad y CSP'],
    ['Flask-Limiter', '-', 'Limitación de peticiones (rate limiting)'],
    ['bcrypt', '-', 'Hash de contraseñas'],
    ['Werkzeug', '-', 'Utilidades y servidor de desarrollo'],
    ['python-docx', '-', 'Generación de documentos Word'],
    ['SQLite', '3.x', 'Base de datos (test.db)'],
    ['HTML5 / CSS3', '-', 'Estructura y estilos del frontend'],
    ['JavaScript / Font Awesome', '-', 'Interactividad e iconografía'],
    ['PWA (manifest / sw.js)', '-', 'Aplicación web progresiva instalable'],
])

doc.add_heading('Beneficios Logrados', level=4)
beneficios = [
    'Automatización de los procesos de inscripción y matrícula escolar.',
    'Trazabilidad total mediante la tabla de auditoría REGISTRO_AUDITORIA.',
    'Reducción significativa de los tiempos de atención a los representantes.',
    'Control de acceso por roles (administrador y secretario) con contraseñas hasheadas.',
    'Publicación de la información institucional en el sitio web público.',
    'Portal del Representante para registro y consulta de estudiantes hijos.',
    'Protección frente a ataques comunes mediante Flask-Talisman y Flask-Limiter.',
    'Disponibilidad de la aplicación como PWA instalable.',
    'Dashboards con métricas clave para la toma de decisiones de la dirección.',
]
for b in beneficios:
    add_bullet(doc, b)

doc.add_heading('Arquitectura del Sistema', level=4)
add_para(doc, 'El sistema fue desarrollado siguiendo una arquitectura de tres capas:')
add_para(doc, 'Capa de Presentación: HTML5, CSS3, JavaScript, Font Awesome, plantillas Flask/Jinja2 '
         'y PWA (manifest.json y sw.js) para la interacción con el usuario.')
add_para(doc, 'Capa de Lógica: Aplicación Flask con vistas (Blueprints), formularios, modelos ORM '
         'SQLAlchemy, autenticación con bcrypt, Flask-CORS, Flask-Talisman y Flask-Limiter.')
add_para(doc, 'Capa de Datos: Base de datos SQLite (test.db), accedida mediante SQLAlchemy 2.0, con '
         '20 tablas que modelan la gestión escolar y el registro de auditoría.')
add_para(doc, 'Despliegue: Railway (https://web-production-d4f80.up.railway.app) con interfaz de usuario '
         'modernizada y responsiva, disponible las 24 horas del día.', bold=True)

doc.add_heading('Estructura de la Base de Datos', level=4)
add_para(doc, 'La base de datos del sistema está compuesta por las siguientes tablas principales:')
db_tables = [
    ['Usuario', 'USUARIO', 'Usuario, contraseña (bcrypt), rol'],
    ['Pais', 'PAIS', 'Países (Venezuela precargado)'],
    ['Estado', 'ESTADO', 'Estados (Distrito Capital)'],
    ['Ciudad', 'CIUDAD', 'Ciudades (Caracas)'],
    ['Representante', 'REPRESENTANTE', 'Datos del representante legal'],
    ['Estudiante', 'ESTUDIANTE', 'Datos del estudiante'],
    ['Familiar', 'FAMILIAR', 'Relación familiar representante-estudiante'],
    ['AnoEscolar', 'ANO_ESCOLAR', 'Período escolar (2025-2026 activo)'],
    ['Grado', 'GRADO', 'Grados 1er a 5to Año'],
    ['Inscripcion', 'INSCRIPCION', 'Matrícula, estado, fecha_retiro, lapso, motivo'],
    ['Noticia', 'NOTICIA', 'Noticias del sitio web'],
    ['ProgramaAcademico', 'PROGRAMA_ACADEMICO', 'Programas (Inicial, Primaria, Media, Especial, Docente)'],
    ['Galeria', 'GALERIA', 'Imágenes de la galería institucional'],
    ['SiteConfig', 'SITE_CONFIG', 'Configuración del sitio web'],
    ['REGISTRO_AUDITORIA', 'REGISTRO_AUDITORIA', 'Log de actividades del sistema'],
]
add_table(doc, ['Entidad', 'Tabla BD', 'Campos Principales'], db_tables)

doc.add_heading('Descripción de los Módulos del Sistema', level=4)
add_para(doc, 'A continuación se describen detalladamente los módulos que componen el sistema:')
modulos_desc = [
    ('Módulo 1: Autenticación / Login',
     'Página de inicio de sesión segura para el personal autorizado. Las contraseñas se almacenan '
     'con hash bcrypt. El acceso está protegido por Flask-Limiter para evitar ataques de fuerza '
     'bruta y por Flask-Talisman que impone cabeceras de seguridad y una Content Security Policy.'),
    ('Módulo 2: Gestión de Usuarios',
     'Permite administrar las cuentas del sistema con dos roles: administrador y secretario. El '
     'usuario administrador puede crear, editar y desactivar cuentas, además de gestionar el año '
     'escolar y la configuración general.'),
    ('Módulo 3: Año Escolar',
     'Gestión de los períodos escolares. Se precarga el Año Escolar 2025-2026 como ACTIVO, sobre el '
     'cual se registran las inscripciones y matrículas.'),
    ('Módulo 4: Estudiantes',
     'Registro, listado y consulta de estudiantes. Incluye datos personales, grado, año escolar y '
     'la asociación con su representante. Soporta búsqueda y filtros.'),
    ('Módulo 5: Representantes',
     'Registro, listado y consulta de los representantes legales. Permite asociar múltiples '
     'estudiantes a un representante mediante la tabla Familiar y consultar la información de sus '
     'hijos.'),
    ('Módulo 6: Matrícula / Inscripción',
     'Gestiona la inscripción de un estudiante a un grado en un año escolar determinado. Registra el '
     'estado de la inscripción, la fecha de retiro, el lapso de registro y el motivo de retiro cuando '
     'aplique, permitiendo llevar el control de la matrícula.'),
    ('Módulo 7: Plantillas de Documentos',
     'Generación de documentos institucionales (constancias, certificados, etc.) utilizando '
     'python-docx a partir de las plantillas definidas.'),
    ('Módulo 8: Administración del Sitio Web',
     'Permite administrar el contenido público: Noticias, Programas Académicos, Galería de imágenes '
     'y la Configuración del sitio (SiteConfig).'),
    ('Módulo 9: Seguridad y Auditoría',
     'Control de acceso basado en roles (administrador, secretario) y registro de todas las '
     'actividades críticas en la tabla REGISTRO_AUDITORIA, lo que garantiza la trazabilidad de los '
     'cambios realizados en el sistema.'),
]
for titulo, desc in modulos_desc:
    doc.add_heading(titulo, level=4)
    add_para(doc, desc)

# ══════════════════════════════════════════════════════════════
# CAPÍTULO IV: CONCLUSIONES Y RECOMENDACIONES
# ══════════════════════════════════════════════════════════════
doc.add_heading('CAPÍTULO IV', level=1)
doc.add_heading('CONCLUSIONES Y RECOMENDACIONES', level=2)

doc.add_heading('4.1. Conclusiones', level=3)
add_para(doc, 'La ejecución de las prácticas profesionales en la U.E. Dr. José Manuel Cova Maza '
           'permitió aplicar los conocimientos adquiridos en la carrera de Ingeniería de Sistemas '
           'para solucionar una necesidad real de la institución: la automatización de sus procesos '
           'de inscripción y administración escolar.')
add_para(doc, 'El sistema INSCRIB SYSTEM, desarrollado con Flask, SQLAlchemy, SQLite y tecnologías '
           'open source, cumple con los objetivos planteados, ofreciendo módulos de gestión de '
           'estudiantes, representantes, matrícula e inscripción, así como un sitio web público '
           'institucional. La incorporación de medidas de seguridad (bcrypt, Flask-Talisman, '
           'Flask-Limiter) y de un registro de auditoría fortalece la integridad y la trazabilidad '
           'de la información.')
add_para(doc, 'El proyecto demostró ser factible técnicamente y socialmente valioso, mejorando la '
           'eficiencia administrativa y la comunicación con la comunidad educativa a través del sitio '
           'web. Asimismo, representó una experiencia de aprendizaje significativa para el pasante, '
           'que integró desarrollo web, bases de datos y seguridad informática.')

doc.add_heading('4.2. Recomendaciones', level=3)
recomendaciones = [
    'Capacitar de forma continua al personal administrativo en el uso del sistema para asegurar su adopción.',
    'Implementar respaldos automatizados periódicos de la base de datos SQLite.',
    'Migrar a un motor de base de datos más robusto (por ejemplo, PostgreSQL) si la institución crece en volumen de datos.',
    'Ampliar el portal del representante con funcionalidades de pagos y seguimiento académico.',
    'Incorporar reportes estadísticos (asistencia, rendimiento) para apoyar la toma de decisiones.',
    'Mantener actualizadas las dependencias (Flask, SQLAlchemy) y las políticas de seguridad del CSP.',
    'El sistema ya se encuentra desplegado en Railway (https://web-production-d4f80.up.railway.app) con interfaz modernizada, disponible las 24 horas.',
]
for i, rec in enumerate(recomendaciones, 1):
    add_para(doc, f'{i}. {rec}')

# ══════════════════════════════════════════════════════════════
# REFERENCIAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('REFERENCIAS', level=1)
referencias = [
    'Pallets. (2024). Flask Documentation. Recuperado de https://flask.palletsprojects.com/',
    'SQLAlchemy. (2024). SQLAlchemy 2.0 Documentation. Recuperado de https://docs.sqlalchemy.org/',
    'Python Software Foundation. (2024). Python 3.13 Documentation. Recuperado de https://docs.python.org/3.13/',
    'python-docx. (2024). python-docx Documentation. Recuperado de https://python-docx.readthedocs.io/',
    'OWASP. (2023). Cheat Sheet Series - Content Security Policy. Recuperado de https://cheatsheetseries.owasp.org/',
    'Universidad Politécnica Territorial "José Antonio Anzoátegui" (UPTJAA). Estructura del Informe de Prácticas Profesionales TSU-3.',
]
for r in referencias:
    add_para(doc, r, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ANEXOS
# ══════════════════════════════════════════════════════════════
doc.add_heading('ANEXOS', level=1)
doc.add_paragraph()
add_para(doc, 'Anexo A: Mapa de Ubicación de la U.E. Dr. José Manuel Cova Maza.')
add_para(doc, 'Anexo B: Organigrama de la Institución.')
add_para(doc, 'Anexo C: Diagrama de Gantt de las 12 semanas de pasantías.')
add_para(doc, 'Anexo D: Bitácora de Actividades Laborales.')
add_para(doc, 'Anexo E: Formato de Evaluación del Pasante.')
add_para(doc, 'Anexo F: Capturas de pantalla del sistema (panel administrativo y sitio público).')
add_para(doc, 'Anexo G: Manual de usuario del sistema INSCRIB SYSTEM.')

doc.save(OUTPUT)
print(f'Documento guardado: {OUTPUT}')
print(f'Tamaño: {os.path.getsize(OUTPUT) / 1024:.1f} KB')
