#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Genera el resumen de avances del proyecto INSCRIB SYSTEM (versión escuela)."""

import os
import datetime
import sqlite3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = r'D:\Usuarios\MCAMPOS\Desktop\INSCRIB SYSTEM\INSCRIB SYSTEM\informe de proyecto'
OUTPUT = os.path.join(OUTPUT_DIR, 'RESUMEN_AVANCES_SGA_v2.docx')
DB_PATH = r'D:\Usuarios\MCAMPOS\Desktop\INSCRIB SYSTEM\INSCRIB SYSTEM\test.db'

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(14)
    hs.paragraph_format.space_after = Pt(6)

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
    return t

# ============================
# PORTADA
# ============================
for _ in range(6):
    add_para('')

add_para('SISTEMA DE GESTIÓN DE INSCRIPCIONES', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20)
add_para('Y ADMINISTRACIÓN ESCOLAR - INSCRIB SYSTEM', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para('')
add_para('RESUMEN COMPLETO DE AVANCES DEL PROYECTO', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para('')
add_para(f'Documento generado el {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")}', align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para('')
add_para('Trayecto 3 - Fase 1 / Pasantías Profesionales', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
add_para('Universidad Politécnica Territorial José Antonio Anzoátegui (UPTJAA)', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
add_para('Marcelo Fernando Campos Anacona', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

doc.add_page_break()

# ============================
# 1. DESCRIPCIÓN GENERAL
# ============================
doc.add_heading('1. Descripción General del Proyecto', level=1)
add_para('El Sistema de Gestión de Inscripciones y Administración Escolar "INSCRIB SYSTEM" es una '
         'aplicación web desarrollada con Flask que gestiona el proceso completo de inscripciones y '
         'matrículas escolares, el registro de estudiantes y representantes, y la administración del '
         'sitio web institucional de la U.E. Dr. José Manuel Cova Maza, institución educativa de '
         'Puerto Ordaz, Estado Bolívar, Venezuela.')
add_para('El sistema reemplaza el método manual de hojas de cálculo (Excel) y registros físicos, '
         'permitiendo un control de matrícula en tiempo real, trazabilidad mediante un registro de '
         'auditoría, generación de documentos, y publicación de información institucional en un '
         'sitio web público accesible desde cualquier navegador.')

# ============================
# 2. ARQUITECTURA TECNOLÓGICA
# ============================
doc.add_heading('2. Arquitectura Tecnológica', level=1)
add_table(
    ['Componente', 'Tecnología', 'Detalle'],
    [
        ['Backend', 'Flask 3.1.1 (Python 3.13)', 'Blueprints, vistas, formularios, context processors'],
        ['ORM', 'Flask-SQLAlchemy 3.1.1 / SQLAlchemy 2.0', 'Mapeo objeto-relacional de las 20 tablas'],
        ['Base de Datos', 'SQLite (test.db)', 'Base de datos local del sistema escolar'],
        ['Seguridad HTTP', 'Flask-Talisman', 'Cabeceras de seguridad y Content Security Policy (CSP)'],
        ['Limitación', 'Flask-Limiter', 'Rate limiting / limitación de peticiones por IP'],
        ['CORS', 'Flask-CORS', 'Políticas de origen cruzado para la API'],
        ['Autenticación', 'bcrypt + Werkzeug', 'Hash de contraseñas y gestión de sesiones'],
        ['Generación Docs', 'python-docx', 'Plantillas de documentos institucionales'],
        ['Frontend', 'HTML5 + CSS3 + JavaScript + Font Awesome', 'Plantillas Jinja2, CSS personalizado'],
        ['PWA', 'manifest.json + sw.js', 'Aplicación web progresiva instalable'],
        ['Servidor', 'Servidor de desarrollo Flask', 'Admin puerto 5001, Sitio público puerto 5002'],
    ]
)

# ============================
# 3. MÓDULOS IMPLEMENTADOS
# ============================
doc.add_heading('3. Módulos del Sistema', level=1)

modulos = [
    ('3.1 Autenticación / Login',
     'Acceso seguro al panel administrativo. Las contraseñas se almacenan con hash bcrypt. El login '
     'está protegido por Flask-Limiter (intentos por IP) y por Flask-Talisman (cabeceras y CSP).',
     [
         ('Iniciar sesión', '/login', 'Formulario de acceso con credenciales hasheadas'),
         ('Cerrar sesión', '/logout', 'Cierre de sesión y limpieza de cookie'),
     ]),
    ('3.2 Gestión de Usuarios',
     'Administración de cuentas con dos roles: administrador y secretario. El administrador puede '
     'crear, editar y desactivar usuarios, y gestionar la configuración general.',
     [
         ('Listar', '/usuarios/', 'Lista paginada de usuarios del sistema'),
         ('Crear', '/usuarios/nuevo/', 'Registro de nuevo usuario con rol'),
         ('Editar', '/usuarios/<id>/editar/', 'Edición de usuario'),
         ('Cambiar contraseña', '/usuarios/<id>/password/', 'Actualización de contraseña (bcrypt)'),
         ('Eliminar', '/usuarios/<id>/eliminar/', 'Eliminación con protección'),
     ]),
    ('3.3 Año Escolar',
     'Gestión de los períodos escolares. Se precarga el Año Escolar 2025-2026 como ACTIVO, sobre el '
     'cual se registran las inscripciones.',
     [
         ('Listar', '/ano-escolar/', 'Lista de años escolares'),
         ('Crear', '/ano-escolar/nuevo/', 'Registro de un período'),
         ('Activar', '/ano-escolar/<id>/activar/', 'Marcar un año como activo'),
         ('Editar', '/ano-escolar/<id>/editar/', 'Edición de período'),
     ]),
    ('3.4 Estudiantes',
     'Registro, listado y consulta de estudiantes. Incluye datos personales, grado, año escolar y '
     'la asociación con su representante.',
     [
         ('Listar', '/estudiantes/', 'Lista paginada con filtros'),
         ('Crear', '/estudiantes/nuevo/', 'Registro de estudiante'),
         ('Detalle', '/estudiantes/<id>/', 'Vista detallada del expediente'),
         ('Editar', '/estudiantes/<id>/editar/', 'Edición de datos'),
         ('Eliminar', '/estudiantes/<id>/eliminar/', 'Eliminación con protección'),
     ]),
    ('3.5 Representantes',
     'Registro, listado y consulta de representantes legales, con asociación a múltiples estudiantes '
     'mediante la tabla Familiar.',
     [
         ('Listar', '/representantes/', 'Lista paginada con filtros'),
         ('Crear', '/representantes/nuevo/', 'Registro de representante'),
         ('Detalle', '/representantes/<id>/', 'Estudiantes asociados'),
         ('Editar', '/representantes/<id>/editar/', 'Edición de datos'),
         ('Eliminar', '/representantes/<id>/eliminar/', 'Eliminación con protección'),
     ]),
    ('3.6 Matrícula / Inscripción',
     'Gestión de la inscripción de un estudiante a un grado en un año escolar. Registra estado, '
     'fecha de retiro, lapso de registro y motivo de retiro.',
     [
         ('Listar', '/inscripciones/', 'Lista de inscripciones por año escolar'),
         ('Crear', '/inscripciones/nueva/', 'Registro de matrícula'),
         ('Detalle', '/inscripciones/<id>/', 'Vista detallada'),
         ('Editar', '/inscripciones/<id>/editar/', 'Edición de matrícula'),
     ]),
    ('3.7 Plantillas de Documentos',
     'Generación de documentos institucionales (constancias, certificados) a partir de plantillas '
     'usando python-docx.',
     [
         ('Generar documento', '/documentos/generar/', 'Generación de documento Word'),
     ]),
    ('3.8 Administración del Sitio Web',
     'Gestión del contenido público: Noticias, Programas Académicos, Galería de imágenes y '
     'Configuración del sitio (SiteConfig).',
     [
         ('Noticias', '/admin/noticias/', 'CRUD de noticias del sitio'),
         ('Programas', '/admin/programas/', 'CRUD de programas académicos'),
         ('Galería', '/admin/galeria/', 'CRUD de imágenes de galería'),
         ('Configuración', '/admin/config/', 'Edición de SiteConfig'),
     ]),
    ('3.9 Seguridad y Auditoría',
     'Control de acceso basado en roles (administrador, secretario) y registro de actividades '
     'críticas en la tabla REGISTRO_AUDITORIA para garantizar trazabilidad.',
     [
         ('Log de auditoría', '/auditoria/', 'Listado de REGISTRO_AUDITORIA'),
         ('Roles', 'admin / secretario', 'Permisos diferenciados por rol'),
     ]),
    ('3.10 Portal del Representante (Sitio Público)',
     'Sitio web institucional público con Inicio, Nosotros, Programas Académicos, Noticias, Galería, '
     'Contacto y Requisitos de Inscripción, además del Portal del Representante para registro y '
     'consulta de estudiantes hijos, perfil y cambio de contraseña.',
     [
         ('Inicio', '/', 'Página principal del sitio'),
         ('Nosotros', '/nosotros/', 'Reseña institucional'),
         ('Programas', '/programas/', 'Programas académicos ofrecidos'),
         ('Noticias', '/noticias/', 'Listado de noticias'),
         ('Galería', '/galeria/', 'Galería de imágenes'),
         ('Contacto', '/contacto/', 'Formulario de contacto'),
         ('Requisitos', '/requisitos/', 'Requisitos de inscripción'),
         ('Portal Representante', '/portal/', 'Registro y consulta de representantes'),
     ]),
]

for titulo, desc, items in modulos:
    doc.add_heading(titulo, level=2)
    add_para(desc)
    add_para('Funcionalidades:', bold=True)
    if isinstance(items, list) and all(isinstance(i, str) for i in items):
        for item in items:
            add_bullet(item)
    elif isinstance(items, list) and all(isinstance(i, (list, tuple)) for i in items):
        add_table(['URL', 'Funcionalidad', 'Descripción'] if len(items[0]) >= 3 else ['URL', 'Descripción'],
                  [i[:3] if len(i) >= 3 else [i[0], i[1], ''] for i in items])
        add_para('')

# ============================
# 4. CORRECCIONES DE BUGS
# ============================
doc.add_heading('4. Correcciones de Bugs y Mejoras', level=1)

add_table(
    ['Bug / Mejora', 'Problema', 'Solución'],
    [
        ['REGISTRO_AUDITORIA', 'La clave primaria no era autoincremental', 'Migración a INTEGER AUTOINCREMENT en create_app'],
        ['INSCRIPCION', 'Faltaban columnas de control de retiro', 'Adición de estado, fecha_retiro, lapso_registro y motivo_retiro'],
        ['Usuarios por defecto', 'No existían credenciales iniciales', 'Creación de admin/admin123 y secretario/secretario123 en create_app'],
        ['Seguridad de contraseñas', 'Contraseñas en texto plano', 'Aplicado hash bcrypt en el almacenamiento'],
        ['Cabeceras HTTP', 'Ausencia de protección', 'Integrado Flask-Talisman con CSP'],
        ['Rate limiting', 'Sin control de peticiones', 'Integrado Flask-Limiter en el login y rutas sensibles'],
        ['CORS', 'Bloqueo de origen cruzado', 'Configurado Flask-CORS para la API'],
        ['PWA', 'Faltaba app instalable', 'Agregados manifest.json y sw.js'],
        ['Documentos', 'Generación manual de constancias', 'Integrado python-docx con plantillas'],
    ]
)

# ============================
# 5. INFRAESTRUCTURA
# ============================
doc.add_heading('5. Infraestructura y Despliegue', level=1)

add_para('5.1 Servidor Local (Flask dev server)', bold=True)
add_bullet('Panel administrativo corriendo en http://localhost:5001')
add_bullet('Sitio web público corriendo en http://localhost:5002')
add_bullet('Base de datos SQLite local (test.db) gestionada por SQLAlchemy')

add_para('5.2 Seguridad', bold=True)
add_bullet('Flask-Talisman: cabeceras de seguridad y Content Security Policy (CSP)')
add_bullet('Flask-Limiter: limitación de peticiones (rate limiting) por IP')
add_bullet('bcrypt: hash de contraseñas de usuarios y representantes')
add_bullet('Roles: administrador y secretario con permisos diferenciados')

add_para('5.3 PWA', bold=True)
add_bullet('manifest.json: metadatos de la aplicación instalable')
add_bullet('sw.js: service worker para soporte offline y caché')

# ============================
# 6. ESTADÍSTICAS (sqlite3 con datos reales)
# ============================
doc.add_heading('6. Estadísticas del Proyecto', level=1)

try:
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
    tables = [r[0] for r in cur.fetchall()]
    counts = []
    for t in tables:
        try:
            cur.execute(f'SELECT COUNT(*) FROM "{t}"')
            c = cur.fetchone()[0]
            counts.append((t, c))
        except Exception:
            pass
    con.close()
    total_tablas = len(tables)
    total_registros = sum(c for _, c in counts)
    add_para(f'El sistema cuenta con {total_tablas} tablas en la base de datos SQLite. A '
             f'continuación se presenta el conteo de registros por tabla (datos reales de test.db):')
    add_table(['Tabla', 'Registros'], sorted(counts, key=lambda x: -x[1]))
    add_para(f'Total de registros en el sistema: {total_registros}.', bold=True)
except Exception as e:
    add_para(f'Nota: No se pudieron obtener estadísticas ({e})')

doc.add_page_break()

# ============================
# 7. PENDIENTES / MEJORAS FUTURAS
# ============================
doc.add_heading('7. Pendientes y Mejoras Futuras', level=1)
add_bullet('Migrar SQLite a un motor de base de datos más robusto (PostgreSQL) para producción.')
add_bullet('Implementar respaldos automatizados programados de la base de datos.')
add_bullet('Ampliar el Portal del Representante con módulo de pagos y seguimiento académico.')
add_bullet('Incorporar reportes estadísticos (matrícula por grado, asistencia, rendimiento).')
add_bullet('Desplegar el sistema en un servidor dedicado con HTTPS (certificado SSL).')
add_bullet('Agregar notificaciones por correo electrónico a representantes.')
add_bullet('Exportar informes a PDF con gráficos (matplotlib + python-docx).')
add_bullet('App móvil nativa o PWA con soporte offline completo.')

doc.add_page_break()

doc.add_heading('8. Resumen de Implementaciones', level=1)

add_para('8.1 Autenticación Segura', bold=True)
add_para('El acceso al panel administrativo utiliza contraseñas hasheadas con bcrypt y está '
         'protegido por Flask-Limiter y Flask-Talisman, fortaleciendo la seguridad frente a ataques '
         'de fuerza bruta y a vulnerabilidades web comunes.')

add_para('8.2 Tabla de Auditoría', bold=True)
add_para('Se implementó la tabla REGISTRO_AUDITORIA con clave primaria INTEGER AUTOINCREMENT que '
         'registra las actividades críticas del sistema, garantizando la trazabilidad de los cambios '
         'realizados por los usuarios.')

add_para('8.3 Control de Matrícula', bold=True)
add_para('El modelo INSCRIPCION fue ampliado con las columnas estado, fecha_retiro, lapso_registro y '
         'motivo_retiro, permitiendo llevar un control detallado de las inscripciones y retiros de '
         'estudiantes durante el año escolar.')

add_para('8.4 Usuarios por Defecto', bold=True)
add_para('En el arranque de la aplicación (create_app) se crean automáticamente el usuario '
         'administrador (admin/admin123) y el usuario secretario (secretario/secretario123), con sus '
         'respectivos roles, facilitando la puesta en marcha inicial.')

add_para('8.5 Sitio Web Público', bold=True)
add_para('Se desarrolló un sitio web institucional con páginas de Inicio, Nosotros, Programas '
         'Académicos, Noticias, Galería, Contacto, Requisitos de Inscripción y el Portal del '
         'Representante, con soporte PWA para instalación como aplicación.')

# ============================
# GUARDAR
# ============================
doc.save(OUTPUT)
print(f'Documento guardado: {OUTPUT}')
print(f'Tamaño: {os.path.getsize(OUTPUT) / 1024:.1f} KB')
