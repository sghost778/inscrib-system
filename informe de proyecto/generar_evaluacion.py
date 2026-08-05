#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Genera el Formato de Evaluación del Pasante - versión INSCRIB SYSTEM (U.E. Dr. José Manuel Cova Maza)."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = r'D:\Usuarios\MCAMPOS\Desktop\INSCRIB SYSTEM\INSCRIB SYSTEM\informe de proyecto'
OUTPUT = os.path.join(OUTPUT_DIR, 'FORMATO_DE_EVALUACION_DEL_PASANTE.docx')

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.2

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.bold = True
    hs.font.size = Pt(16 - level * 2)

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_field_table(doc, left, right):
    t = doc.add_table(rows=1, cols=2)
    t.columns[0].width = Cm(5)
    t.columns[1].width = Cm(12)
    c0 = t.cell(0, 0).paragraphs[0]
    r0 = c0.add_run(left); r0.bold = True; r0.font.name = 'Times New Roman'; r0.font.size = Pt(11)
    c1 = t.cell(0, 1).paragraphs[0]
    r1 = c1.add_run(right); r1.font.name = 'Times New Roman'; r1.font.size = Pt(11)
    return t

def bordered_table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val)); r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return t

# ============ EVALUACIÓN TUTOR INDUSTRIAL (0-2) ============
doc.add_heading('EVALUACIÓN DEL DESEMPEÑO DEL (DE LA) ESTUDIANTE POR', level=1)
add_para('TUTOR(A) INDUSTRIAL', bold=True)
add_para('LICENCIADO-INGENIERÍA', bold=True)
doc.add_paragraph()

add_field_table(doc, 'Apellido y Nombres:', 'Marcelo Fernando Campos Anacona')
add_field_table(doc, 'Cédula de Identidad:', 'V-32.062.637')
add_field_table(doc, 'PNF que cursa:', 'Ingeniería de Sistemas (UPTJAA)')
add_field_table(doc, 'Nombre de la institución:', 'U.E. Dr. José Manuel Cova Maza')
add_field_table(doc, 'Departamento donde se efectuó la Práctica Profesional:', 'Departamento de Administración Escolar / Informática')
add_field_table(doc, 'Tutor(a) Institucional:', 'Sol Liliana Romero Gil (Coordinadora Académica)')
add_field_table(doc, 'Fechas de la Práctica Profesional:', 'Inicio: 01 de Junio de 2026  |  Culminación: 21 de Agosto de 2026')

doc.add_paragraph()
aspectos_0_2 = [
    'Cumplimiento del horario normal de trabajo',
    'Capacidad para proponer espontánea y oportunamente sugerencias útiles para la organización,',
    'Facilidad de comunicación verbal y escrita; habilidad para dar a conocer y defender sus ideas.',
    'Receptividad a planteamientos diferentes a los presentados por él.',
    'Responsabilidad y puntualidad en la ejecución de las actividades cumpliendo con las condiciones de tiempo y calidad preestablecidas.',
    'Cumplimiento y aplicación de normas de seguridad y de prevención de accidentes.',
    'Disposición para colaborar con los (las) compañeros (as) de trabajo y con los (las) supervisores(as) en forma permanente y espontánea.',
    'Adaptación a situaciones cambiantes o demandas del entorno.',
    'Productividad en función de las metas planificadas y alcanzadas durante el periodo de prácticas.',
    'Calidad de los resultados que presenta como producto de su trabajo.',
    'Manejo y conocimientos de técnicas y procedimientos inherentes a las actividades asignadas.',
    'Habilidad y destreza en el manejo de herramientas informáticas para la solución de problemas.',
    'Compromiso con las metas de la empresa u organización.',
    'Habilidades para establecer relaciones interpersonales y facilidad para el trabajo en equipo.',
    'Capacidad de aprendizaje rápido, para el Trabajo en equipo y bajo presión',
]
rows = [[str(i + 1), a, '0-2', '0-2', ''] for i, a in enumerate(aspectos_0_2)]
rows.append(['Calificación Final (Sumatoria de todos los Aspectos Evaluados)', '', '', '', '_____ 30'])
bordered_table(['Ítems', 'Aspectos Evaluados', 'Intervalo de Ponderación', 'Intervalo de Calificación', 'Calificación parcial'], rows)
doc.add_page_break()

# ============ EVALUACIÓN TUTOR INDUSTRIAL (0-3) ============
doc.add_heading('EVALUACIÓN DEL DESEMPEÑO DEL (DE LA) ESTUDIANTE POR', level=1)
add_para('TUTOR(A) INDUSTRIAL', bold=True)
add_para('TSU', bold=True)
doc.add_paragraph()

add_field_table(doc, 'Apellido y Nombres:', 'Marcelo Fernando Campos Anacona')
add_field_table(doc, 'Cédula de Identidad:', 'V-32.062.637')
add_field_table(doc, 'PNF que cursa:', 'Ingeniería de Sistemas (UPTJAA)')
add_field_table(doc, 'Nombre de la institución:', 'U.E. Dr. José Manuel Cova Maza')
add_field_table(doc, 'Departamento donde se efectuó la Práctica Profesional:', 'Departamento de Administración Escolar / Informática')
add_field_table(doc, 'Tutor(a) Institucional:', 'Sol Liliana Romero Gil (Coordinadora Académica)')
add_field_table(doc, 'Fechas de la Práctica Profesional:', 'Inicio: 01 de Junio de 2026  |  Culminación: 21 de Agosto de 2026')

doc.add_paragraph()
aspectos_0_3 = [
    'Cumplimiento del horario normal de trabajo',
    'Capacidad para proponer espontánea y oportunamente sugerencias útiles para la organización,',
    'Facilidad de comunicación verbal y escrita; habilidad para dar a conocer y defender sus ideas.',
    'Responsabilidad, Puntualidad Y Personalidad.',
    'Capacidad Y Disposición Para Relacionarse Con Los Miembros De La Organización.',
    'Dominio De Conocimientos Básicos Relacionados Con Las Tareas A Ejecutar.',
    'Capacidad Para Detectar Problemas. Iniciativa Para Ayudar A Resolver Problemas.',
    'Destrezas En La Realización De Las Tareas Asignadas.',
    'Habilidades para establecer relaciones interpersonales y facilidad para el trabajo en equipo.',
    'Capacidad de aprendizaje rápido, para el Trabajo en equipo y bajo presión',
]
rows2 = [[str(i + 1), a, '0-3', '0-3', ''] for i, a in enumerate(aspectos_0_3)]
rows2.append(['Calificación Final (Sumatoria de todos los Aspectos Evaluados)', '', '', '', '_____ 30'])
bordered_table(['Ítems', 'Aspectos Evaluados', 'Intervalo de Ponderación', 'Intervalo de Calificación', 'Calificación parcial'], rows2)
doc.add_page_break()

# ============ EVALUACIÓN TUTOR ACADÉMICO ============
doc.add_heading('EVALUACIÓN DEL (DE LA) ESTUDIANTE POR TUTOR(A) ACADÉMICO', level=1)
doc.add_paragraph()
add_field_table(doc, 'Apellido y Nombres:', 'Marcelo Fernando Campos Anacona')
add_field_table(doc, 'Cédula de Identidad:', 'V-32.062.637')
add_field_table(doc, 'Carrera:', 'Ingeniería de Sistemas (UPTJAA)')
add_field_table(doc, 'Nombre de la institución:', 'U.E. Dr. José Manuel Cova Maza')
add_field_table(doc, 'Tutor(a) Académico:', 'Ing. Yalitza Guevara')
add_field_table(doc, 'Fechas de la Práctica Profesional:', 'Inicio: 01 de Junio de 2026  |  Culminación: 21 de Agosto de 2026')
doc.add_paragraph()
add_para('El Tutor Académico evaluará el informe de pasantías y el desempeño del estudiante según '
         'el formato institucional de la UPTJAA, consignando la calificación correspondiente en la '
         'casilla habilitada a continuación:')
add_para('Calificación final del Tutor Académico:  _____', bold=True)

# ============ NOTA SOBRE DESPLIEGUE ============
doc.add_heading('NOTA SOBRE DESPLIEGUE DEL SISTEMA', level=1)
doc.add_paragraph()
add_para(doc, 'El sistema INSCRIB SYSTEM fue desplegado exitosamente en la plataforma Railway '
         '(https://web-production-d4f80.up.railway.app), disponible las 24 horas del día para '
         'el personal autorizado. El despliegue incluye:')
add_para(doc, '• Panel administrativo (admin y secretario)')
add_para(doc, '• Sitio web público institucional')
add_para(doc, '• Portal del Representante')
add_para(doc, '• Interfaz de usuario modernizada y responsiva')
add_para(doc, '• Seguridad reforzada (Flask-Talisman, Flask-Limiter, bcrypt)')

doc.save(OUTPUT)
print(f'Documento guardado: {OUTPUT}')
print(f'Tamaño: {os.path.getsize(OUTPUT) / 1024:.1f} KB')
