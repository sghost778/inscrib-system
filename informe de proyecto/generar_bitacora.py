#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Genera la Bitácora de Actividades Laborales de Pasantes - versión INSCRIB SYSTEM."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = r'D:\Usuarios\MCAMPOS\Desktop\INSCRIB SYSTEM\INSCRIB SYSTEM\informe de proyecto'
OUTPUT = os.path.join(OUTPUT_DIR, 'BITACORA_ACTIVIDADES.docx')

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.bold = True
    hs.font.size = Pt(16 - level * 2)

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_field_table(doc, left, right):
    t = doc.add_table(rows=1, cols=2)
    t.columns[0].width = Cm(5)
    t.columns[1].width = Cm(12)
    c0 = t.cell(0, 0).paragraphs[0]
    r0 = c0.add_run(left); r0.bold = True; r0.font.name = 'Times New Roman'; r0.font.size = Pt(12)
    c1 = t.cell(0, 1).paragraphs[0]
    r1 = c1.add_run(right); r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
    for row in t.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            b = OxmlElement('w:tcBorders')
            for bn in ('top', 'left', 'bottom', 'right'):
                e = OxmlElement(f'w:{bn}')
                e.set(qn('w:val'), 'none'); e.set(qn('w:sz'), '0')
                e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'auto')
                b.append(e)
            tcPr.append(b)
    return t

def add_week_table(title, rows):
    doc.add_heading(title, level=3)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(['Día / Fecha', 'Actividades Realizadas']):
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    for dia, act in rows:
        cells = t.add_row().cells
        p0 = cells[0].paragraphs[0]; r0 = p0.add_run(dia); r0.font.name = 'Times New Roman'; r0.font.size = Pt(11)
        p1 = cells[1].paragraphs[0]; r1 = p1.add_run(act); r1.font.name = 'Times New Roman'; r1.font.size = Pt(11)

# Título
doc.add_heading('BITÁCORA DE ACTIVIDADES LABORALES DE PASANTES', level=1)

# Datos de la empresa
doc.add_heading('DATOS DE LA ORGANIZACIÓN', level=2)
add_field_table(doc, 'Nombre:', 'U.E. Dr. José Manuel Cova Maza')
add_field_table(doc, 'Dirección:', 'Av. Principal, Puerto Ordaz, Estado Bolívar')
add_field_table(doc, 'Teléfono:', '+58 412-1234567')
add_field_table(doc, 'Correo electrónico:', 'info@uejmcm.edu.ve')
add_field_table(doc, 'Sector:', 'Educación')

# Datos del pasante
doc.add_heading('DATOS DEL PASANTE', level=2)
add_field_table(doc, 'Nombre:', 'Marcelo Fernando Campos Anacona')
add_field_table(doc, 'C.I.:', 'V-32.062.637')
add_field_table(doc, 'Carrera:', 'Ingeniería de Sistemas - UPTJAA')
add_field_table(doc, 'Nivel:', 'Trayecto 3, Fase 1')
add_field_table(doc, 'Período:', '01 de Junio al 21 de Agosto de 2026')

# Datos del tutor industrial
doc.add_heading('DATOS DEL TUTOR INDUSTRIAL', level=2)
add_field_table(doc, 'Nombre:', 'Sol Liliana Romero Gil')
add_field_table(doc, 'Cargo:', 'Coordinadora Académica')
add_field_table(doc, 'Correo:', 'info@uejmcm.edu.ve')
add_field_table(doc, 'Teléfono:', '+58 412-1234567')

# Datos del tutor académico
doc.add_heading('DATOS DEL TUTOR ACADÉMICO', level=2)
add_field_table(doc, 'Nombre:', 'Ing. Yalitza Guevara')
add_field_table(doc, 'C.I.:', 'V-13.498.102')

semanas = [
    ('SEMANA N° 1 - (01 al 05 de Junio 2026)', [
        ('Lunes 01/06', 'Presentación a la tutora industrial y al personal involucrado en la pasantía. Recorrido por las instalaciones de la institución.'),
        ('Martes 02/06', 'Participación en charlas de higiene, seguridad y ambiente (PSA) de la institución. Entrega de credenciales.'),
        ('Miércoles 03/06', 'Inducción al Departamento de Administración Escolar / Informática. Conocimiento de los procesos de inscripción.'),
        ('Jueves 04/06', 'Revisión de procedimientos administrativos de la institución. Documentación de procesos actuales.'),
        ('Viernes 05/06', 'Asignación del tema de pasantías por parte de la tutora. Definición del alcance del proyecto INSCRIB SYSTEM.'),
    ]),
    ('SEMANA N° 2 - (08 al 12 de Junio 2026)', [
        ('Lunes 08/06', 'Elaboración del plan de trabajo en conjunto con la tutora industrial.'),
        ('Martes 09/06', 'Culminación del plan de trabajo. Definición de cronograma y actividades.'),
        ('Miércoles 10/06', 'Revisión de las hojas de cálculo usadas para el control de estudiantes y matrículas.'),
        ('Jueves 11/06', 'Inducción sobre los niveles educativos ofrecidos (Inicial, Primaria, Media, Especial, Docente).'),
        ('Viernes 12/06', 'Identificación de las necesidades de información de los representantes. Documentación fotográfica.'),
    ]),
    ('SEMANA N° 3 - (15 al 19 de Junio 2026)', [
        ('Lunes 15/06', 'Análisis de los procesos de inscripción y matrícula escolar.'),
        ('Martes 16/06', 'Levantamiento de requerimientos funcionales del sistema.'),
        ('Miércoles 17/06', 'Análisis del flujo de información entre Secretaría, Coordinación y Docencia.'),
        ('Jueves 18/06', 'Identificación de los actores (administrador, secretario, representante).'),
        ('Viernes 19/06', 'Documentación de especificaciones para el sistema INSCRIB SYSTEM.'),
    ]),
    ('SEMANA N° 4 - (22 al 26 de Junio 2026)', [
        ('Lunes 22/06', 'Diseño de la arquitectura del sistema con Flask y SQLAlchemy.'),
        ('Martes 23/06', 'Selección de tecnologías (Flask-CORS, Flask-Talisman, Flask-Limiter, bcrypt, PWA).'),
        ('Miércoles 24/06', 'Diseño de la base de datos y modelamiento de entidades.'),
        ('Jueves 25/06', 'Definición de roles y política de seguridad.'),
        ('Viernes 26/06', 'Configuración del entorno de desarrollo y del servidor Flask.'),
    ]),
    ('SEMANA N° 5 - (29 de Junio al 03 de Julio 2026)', [
        ('Lunes 29/06', 'Desarrollo del módulo de Autenticación (login seguro con bcrypt).'),
        ('Martes 30/06', 'Implementación del modelo de datos para Año Escolar y Grados.'),
        ('Miércoles 01/07', 'Desarrollo de la gestión de Usuarios con roles (admin y secretario).'),
        ('Jueves 02/07', 'Configuración de Flask-Talisman (cabeceras de seguridad y CSP).'),
        ('Viernes 03/07', 'Pruebas de autenticación y control de acceso.'),
    ]),
    ('SEMANA N° 6 - (06 al 10 de Julio 2026)', [
        ('Lunes 06/07', 'Desarrollo del módulo de Estudiantes (registro y listado).'),
        ('Martes 07/07', 'Implementación de formularios de registro con validaciones.'),
        ('Miércoles 08/07', 'Desarrollo de la consulta de estudiantes por año escolar y grado.'),
        ('Jueves 09/07', 'Integración con el módulo de Representantes.'),
        ('Viernes 10/07', 'Pruebas del módulo de estudiantes y ajustes.'),
    ]),
    ('SEMANA N° 7 - (13 al 17 de Julio 2026)', [
        ('Lunes 13/07', 'Desarrollo del módulo de Representantes (registro y listado).'),
        ('Martes 14/07', 'Implementación de la relación familiar Representante-Estudiante (tabla Familiar).'),
        ('Miércoles 15/07', 'Desarrollo de la consulta de representantes y estudiantes asociados.'),
        ('Jueves 16/07', 'Generación de documentos mediante plantillas (python-docx).'),
        ('Viernes 17/07', 'Pruebas del módulo de representantes.'),
    ]),
    ('SEMANA N° 8 - (20 al 24 de Julio 2026)', [
        ('Lunes 20/07', 'Desarrollo del módulo de Matrícula/Inscripción.'),
        ('Martes 21/07', 'Implementación de columnas de control (estado, fecha_retiro, lapso_registro, motivo_retiro).'),
        ('Miércoles 22/07', 'Desarrollo de la asociación inscripción-estudiante-año escolar-grado.'),
        ('Jueves 23/07', 'Implementación de reportes de matrícula por grado.'),
        ('Viernes 24/07', 'Pruebas del módulo de matrícula e inscripción.'),
    ]),
    ('SEMANA N° 9 - (27 al 31 de Julio 2026)', [
        ('Lunes 27/07', 'Desarrollo del sitio web público (Inicio, Nosotros, Programas).'),
        ('Martes 28/07', 'Implementación de Noticias, Galería y Formulario de Contacto.'),
        ('Miércoles 29/07', 'Desarrollo del Portal del Representante (registro y consulta).'),
        ('Jueves 30/07', 'Desarrollo del módulo de Seguridad y Auditoría (REGISTRO_AUDITORIA).'),
        ('Viernes 31/07', 'Pruebas integrales del sistema.'),
    ]),
    ('SEMANA N° 10 - (03 al 07 de Agosto 2026)', [
        ('Lunes 03/08', 'Implementación de Flask-Limiter (limitación de peticiones).'),
        ('Martes 04/08', 'Pruebas de funcionamiento del sistema completo (admin y sitio público).'),
        ('Miércoles 05/08', 'Corrección de errores identificados en las pruebas.'),
        ('Jueves 06/08', 'Optimización de consultas y ajustes de la interfaz de usuario.'),
        ('Viernes 07/08', 'Documentación técnica del sistema.'),
    ]),
    ('SEMANA N° 11 - (10 al 14 de Agosto 2026)', [
        ('Lunes 10/08', 'Puesta en marcha del sistema (admin puerto 5001, sitio público puerto 5002).'),
        ('Martes 11/08', 'Capacitación al personal de la Secretaría y Coordinación Académica.'),
        ('Miércoles 12/08', 'Capacitación a usuarios del sistema (demostración práctica).'),
        ('Jueves 13/08', 'Resolución de dudas y soporte técnico post-implementación.'),
        ('Viernes 14/08', 'Documentación de usuario y manuales.'),
    ]),
    ('SEMANA N° 12 - (17 al 21 de Agosto 2026)', [
        ('Lunes 17/08', 'Revisión final del sistema. Ajustes menores.'),
        ('Martes 18/08', 'Elaboración del informe de pasantías.'),
        ('Miércoles 19/08', 'Continuación del informe. Inclusión de resultados.'),
        ('Jueves 20/08', 'Revisión y corrección del informe.'),
        ('Viernes 21/08', 'Entrega del informe y cierre de pasantías.'),
    ]),
]

for title, rows in semanas:
    add_week_table(title, rows)

doc.save(OUTPUT)
print(f'Documento guardado: {OUTPUT}')
print(f'Tamaño: {os.path.getsize(OUTPUT) / 1024:.1f} KB')
