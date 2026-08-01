from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

title = doc.add_heading('INSCRIB SYSTEM - Sistema de Gestion Escolar', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
doc.add_paragraph('Documento de Resumen del Proyecto', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Escuela: Escuela José Manuel Cova Maza', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

doc.add_heading('1. Descripcion General', level=1)
doc.add_paragraph(
    'INSCRIB SYSTEM es un sistema de gestion escolar integral desarrollado en Flask (Python) '
    'con base de datos SQLite. El sistema cuenta con dos componentes principales:'
)

items = [
    ('Sitio Web Publico', 'Pagina informativa visible para todo publico sin necesidad de iniciar sesion.'),
    ('Panel Administrativo', 'Area privada con login para gestionar estudiantes, representantes, matricula y contenido del sitio web.')
]
for bold, text in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold + ': ')
    run.bold = True
    p.add_run(text)

doc.add_heading('2. Sitio Web Publico', level=1)
doc.add_paragraph('El sitio web publico se accede desde la raiz (/) y contiene las siguientes secciones:')

sections = [
    ('Inicio (/)', 'Pagina principal con hero banner, "Por que elegirnos", ultimas noticias y formulario de contacto.'),
    ('Nosotros (/nosotros)', 'Informacion institucional: mision, vision, valores.'),
    ('Programas (/programas)', 'Programas academicos cargados desde BD (Inicial, Primaria, Bachillerato).'),
    ('Noticias (/noticias)', 'Noticias y eventos escolares. Cada noticia tiene pagina de detalle en /noticia/<id>.'),
    ('Detalle de Noticia (/noticia/<id>)', 'Pagina individual con contenido completo de cada noticia.'),
    ('Galeria (/galeria)', 'Album de imagenes con vista previa y modal para ampliar.'),
    ('Contacto (/contacto)', 'Formulario de contacto publico que envia mensajes a la BD.')
]
for t, desc in sections:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3. Panel Administrativo (Gestion Interna)', level=1)
doc.add_paragraph(
    'Login en /login. Usuarios por defecto: admin/admin123 (admin), secretario/secretario123 (secretario). '
    'Soporta multiples roles (admin, secretario, docente). Menu lateral con opciones:'
)

admin_sections = [
    'Inicio - Dashboard con ultimos accesos',
    'Estudiantes - Listado, consulta y registro',
    'Representantes - Listado, consulta y registro',
    'Plantilla de Inscripcion - Registro de matricula',
    'Matricula - Control de inscripciones activas',
    'Gestion de Ano - Apertura/cierre de anos escolares',
    'Usuarios - CRUD completo de usuarios del sistema',
    'Configurar Sitio Web - Admin completo del contenido publico'
]
for s in admin_sections:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('4. Configuracion del Sitio Web', level=1)
doc.add_paragraph('Panel completo para administrar el contenido del sitio publico:')

config_items = [
    ('Configuracion General', 'Nombre, telefono, email, direccion y horario de la escuela.'),
    ('Noticias', 'Crear, editar y eliminar noticias. Soporta subida de imagenes.'),
    ('Programas Academicos', 'CRUD de programas educativos con nombre, descripcion, nivel e icono.'),
    ('Galeria', 'Agregar y eliminar imagenes con subida de archivos.'),
    ('Mensajes', 'Visualizar mensajes del formulario de contacto.')
]
for t, desc in config_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('5. Roles de Usuario', level=1)
doc.add_paragraph('El sistema soporta 3 roles de usuario:')
roles = [
    ('admin', 'Acceso completo a todas las funciones del sistema.'),
    ('secretario', 'Gestion de estudiantes, representantes e inscripciones.'),
    ('docente', 'Acceso limitado a consulta de estudiantes y matricula.')
]
for r, desc in roles:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(r + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('6. Sistema de Subida de Archivos', level=1)
doc.add_paragraph(
    'Endpoint /api/upload permite subir imagenes al servidor. Las imagenes se guardan '
    'en /static/uploads/ con nombres unicos UUID. Formatos permitidos: png, jpg, jpeg, gif, webp, svg.'
)

doc.add_heading('7. Tecnologias Utilizadas', level=1)
techs = [
    'Backend: Python 3 + Flask',
    'Base de Datos: SQLite con SQLAlchemy ORM',
    'Frontend: HTML5, CSS3, JavaScript (vanilla)',
    'Estilos: Font Awesome 6, diseno responsivo',
    'Seguridad: bcrypt, flask-limiter, flask-talisman',
    'Subida de Archivos: Werkzeug secure_filename + UUID'
]
for t in techs:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('8. Estructura de la Base de Datos', level=1)
doc.add_paragraph('El sistema cuenta con 20 tablas:')

tables = [
    'USUARIO - Usuarios con soporte de roles (admin/secretario/docente)',
    'ESTUDIANTE - Datos de los estudiantes',
    'REPRESENTANTE - Datos de los representantes',
    'INSCRIPCION - Registro de matricula',
    'GRADO - Grados academicos',
    'ANO_ESCOLAR - Anos escolares',
    'FAMILIAR - Datos de padres/madres',
    'PAIS, ESTADO, CIUDAD - Datos geograficos',
    'REGISTRO_AUDITORIA - Registro de actividades',
    'SITE_CONFIG - Configuracion del sitio web',
    'NOTICIA - Noticias y eventos',
    'PROGRAMA_ACADEMICO - Programas educativos',
    'GALERIA - Imagenes de la galeria',
    'MENSAJE_CONTACTO - Mensajes de contacto'
]
for t in tables:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('9. Mejoras y Correcciones Aplicadas', level=1)
fixes = [
    'Corregido bug: Grado.capacidad no existia en el modelo',
    'Corregido bug: Grado.nivel tenia NOT NULL sin valor por defecto',
    'Agregado: Pagina de detalle de noticia (/noticia/<id>)',
    'Agregado: Sistema de subida de imagenes con /api/upload',
    'Agregado: Roles de usuario (admin, secretario, docente)',
    'Agregado: CRUD completo de usuarios via API',
    'Agregado: Sidebar responsive (colapsable en movil)',
    'Agregado: Datos de semilla para noticias, programas y galeria',
    'Agregado: Usuario secundario "secretario" para pruebas',
    'Agregado: requirements.txt con dependencias del proyecto',
    'Corregido: UnicodeEncodeError en prints con emojis',
    'Actualizado: Login devuelve rol del usuario',
    'Actualizado: Todos los enlaces internos (/, /login, etc.)'
]
for f in fixes:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('10. Como Ejecutar el Sistema', level=1)
steps = [
    'Abrir terminal en la carpeta del proyecto',
    'pip install -r requirements.txt (solo la primera vez)',
    'python create_db_table.py (crea la BD con datos de prueba)',
    'python app.py',
    'Abrir navegador en http://localhost:5000 (sitio publico)',
    'Ir a http://localhost:5000/login (panel admin)',
    'Usuarios: admin/admin123 o secretario/secretario123'
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph('')
doc.add_paragraph('--- Fin del Documento ---').alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'RESUMEN_PROYECTO_INSCRIB.docx')
doc.save(output_path)
print(f'Documento creado: {output_path}')
